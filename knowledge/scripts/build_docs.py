#!/usr/bin/env python3
"""
knowledge/ の Markdown 一式を 1 冊にまとめて HTML / PDF / Word(.docx) を生成する。
  実行: /Users/shimadakoutaro/shoken/.venv/bin/python knowledge/scripts/build_docs.py
  依存: markdown, python-docx, playwright（PDF）— shoken の venv に入っている
出力: knowledge/dist/oshima-tourism-ai-knowledge.{html,pdf,docx}
"""
import re, datetime as dt
from pathlib import Path

import markdown
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
DIST = ROOT / "dist"
TITLE = "伊豆大島 観光AIナレッジ — 情報基盤ドキュメント"
SUB = "ChatGPTに時刻入りの旅程を提案させるための、AI-readableな時刻表と現地の暗黙知（SBC.別班連携チーム）"

ORDER = [
    ("はじめに", ROOT / "README.md"),
    ("決定事項の記録", ROOT / "docs/DECISIONS.md"),
    ("ライセンス・コンプライアンス", ROOT / "docs/LICENSE-COMPLIANCE.md"),
    ("ODPT登録とGTFS取得", ROOT / "docs/ODPT-SETUP.md"),
    ("ChatGPTへの渡し方", ROOT / "docs/HOW-TO-USE-WITH-CHATGPT.md"),
]
KNOW = sorted((ROOT / "oshima/knowledge").glob("*.md"))

def load(p):
    t = p.read_text(encoding="utf-8")
    # 内部リンク [text](xxx.md) → text（結合文書では不要）
    t = re.sub(r"\[([^\]]+)\]\((?!http)[^)]+\.md[^)]*\)", r"\1", t)
    return t

def combined_md():
    parts = [f"# {TITLE}\n\n{SUB}\n\n生成日：{dt.date.today().isoformat()}　ブランチ：feature/ai-readable-knowledge\n"]
    for label, p in ORDER:
        body = load(p)
        # 先頭の H1 をそのまま章題に、以降の見出しを1段下げる
        body = re.sub(r"^(#{1,5}) ", lambda m: "#" + m.group(1) + " ", body, flags=re.M)
        parts.append(f"\n\n# 第{len(parts)}章 {label}\n\n" + body)
    kparts = []
    for p in KNOW:
        b = load(p)
        b = re.sub(r"^(#{1,5}) ", lambda m: "#" + m.group(1) + " ", b, flags=re.M)
        kparts.append(b)
    parts.append(f"\n\n# 第{len(parts)}章 現地の暗黙知（AI-readable）\n\n" + "\n\n".join(kparts))
    return "\n".join(parts)

CSS = """
:root{--sea:#0b5aa6;--sea-deep:#083c70;--sea-light:#e6f0fb;--tsubaki:#d3283b;--line:#dde3ea;--sub:#5b6773}
*{box-sizing:border-box}body{margin:0;font-family:-apple-system,BlinkMacSystemFont,"Hiragino Sans","Hiragino Kaku Gothic ProN","Yu Gothic",Meiryo,sans-serif;color:#1f2933;line-height:1.75;font-size:15px;background:#fff}
.wrap{max-width:860px;margin:0 auto;padding:28px 24px 60px}
h1{font-size:1.7rem;color:var(--sea-deep);border-bottom:3px solid var(--tsubaki);padding-bottom:6px;margin:44px 0 16px;page-break-before:always}
h1:first-of-type{page-break-before:auto;font-size:2rem;border:0;margin-top:0}
h2{font-size:1.25rem;color:var(--sea-deep);border-left:6px solid var(--tsubaki);padding-left:10px;margin:30px 0 10px}
h3{font-size:1.05rem;color:var(--sea);margin:22px 0 8px}
table{border-collapse:collapse;width:100%;font-size:.9rem;margin:10px 0}th,td{border:1px solid var(--line);padding:6px 9px;vertical-align:top;text-align:left}th{background:var(--sea-light);color:var(--sea-deep)}
code{font-family:ui-monospace,Menlo,Consolas,monospace;background:#eef1f5;padding:1px 5px;border-radius:4px;font-size:.88em}
pre{background:#0f1c2b;color:#dbe7f5;padding:12px;border-radius:8px;overflow-x:auto;font-size:.82rem;line-height:1.5}pre code{background:none;color:inherit;padding:0}
blockquote{margin:10px 0;padding:8px 14px;background:#fff6dc;border-left:4px solid #f2a900;border-radius:4px;color:#3a2a00}
a{color:var(--sea)}.toc{background:var(--sea-light);border-radius:10px;padding:12px 18px}.toc li{margin:2px 0}
input[type=checkbox]{margin-right:6px}
@media print{.wrap{max-width:none;padding:0}h1{page-break-before:always}h1:first-of-type{page-break-before:auto}pre,table,blockquote{page-break-inside:avoid}}
"""

def build_html(md_text):
    html = markdown.markdown(md_text, extensions=["tables", "fenced_code", "sane_lists", "toc"])
    html = html.replace("[ ]", '<input type="checkbox" disabled>').replace("[x]", '<input type="checkbox" checked disabled>')
    doc = f"""<!DOCTYPE html><html lang="ja"><head><meta charset="UTF-8"><meta name="viewport" content="width=device-width, initial-scale=1"><title>{TITLE}</title><style>{CSS}</style></head><body><div class="wrap">{html}</div></body></html>"""
    (DIST / "oshima-tourism-ai-knowledge.html").write_text(doc, encoding="utf-8")
    return DIST / "oshima-tourism-ai-knowledge.html"

def build_pdf(html_path):
    from playwright.sync_api import sync_playwright
    with sync_playwright() as p:
        b = p.chromium.launch(); pg = b.new_page()
        pg.goto(f"file://{html_path}", wait_until="networkidle"); pg.wait_for_timeout(500)
        pg.pdf(path=str(DIST / "oshima-tourism-ai-knowledge.pdf"), format="A4", print_background=True,
               margin={"top": "16mm", "bottom": "16mm", "left": "14mm", "right": "14mm"},
               display_header_footer=True, header_template="<div></div>",
               footer_template='<div style="font-size:9px;color:#888;width:100%;text-align:center;font-family:sans-serif">伊豆大島 観光AIナレッジ ／ <span class="pageNumber"></span> / <span class="totalPages"></span></div>')
        b.close()

# ---------- Markdown → docx（必要十分なサブセット） ----------
def add_runs(par, text):
    """**bold**, `code`, [text](url) を run に分解"""
    pos = 0
    for m in re.finditer(r"(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\))", text):
        if m.start() > pos: par.add_run(text[pos:m.start()])
        tok = m.group(0)
        if tok.startswith("**"): r = par.add_run(tok[2:-2]); r.bold = True
        elif tok.startswith("`"): r = par.add_run(tok[1:-1]); r.font.name = "Menlo"; r.font.size = Pt(9.5)
        else:
            mm = re.match(r"\[([^\]]+)\]\(([^)]+)\)", tok); r = par.add_run(mm.group(1))
            r.font.color.rgb = RGBColor(0x0b, 0x5a, 0xa6); r.underline = True
            par.add_run(f"（{mm.group(2)}）").font.size = Pt(8)
        pos = m.end()
    if pos < len(text): par.add_run(text[pos:])

def set_jp_font(doc):
    st = doc.styles["Normal"]; st.font.name = "Hiragino Sans"; st.font.size = Pt(10.5)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), "Hiragino Sans")
    for name in ["Heading 1", "Heading 2", "Heading 3", "Heading 4", "Title"]:
        s = doc.styles[name]; s.font.name = "Hiragino Sans"; s.element.rPr.rFonts.set(qn("w:eastAsia"), "Hiragino Sans")
        s.font.color.rgb = RGBColor(0x08, 0x3c, 0x70)

def build_docx(md_text):
    doc = Document(); set_jp_font(doc)
    for s in doc.sections: s.left_margin = s.right_margin = Cm(2); s.top_margin = s.bottom_margin = Cm(2)
    lines = md_text.split("\n"); i = 0; first_h1 = True
    while i < len(lines):
        ln = lines[i]
        if ln.startswith("```"):
            buf = []; i += 1
            while i < len(lines) and not lines[i].startswith("```"): buf.append(lines[i]); i += 1
            p = doc.add_paragraph(); r = p.add_run("\n".join(buf)); r.font.name = "Menlo"; r.font.size = Pt(8.5)
            p.paragraph_format.left_indent = Cm(0.5); i += 1; continue
        m = re.match(r"^(#{1,6})\s+(.*)", ln)
        if m:
            lvl = len(m.group(1)); text = m.group(2)
            if lvl == 1 and first_h1:
                doc.add_heading(text, 0); first_h1 = False
            else:
                if lvl == 1: doc.add_page_break()
                doc.add_heading(text, min(lvl, 4))
            i += 1; continue
        if ln.startswith("|") and i + 1 < len(lines) and re.match(r"^\|?\s*:?-{2,}", lines[i + 1]):
            hdr = [c.strip() for c in ln.strip().strip("|").split("|")]
            rows = []; i += 2
            while i < len(lines) and lines[i].startswith("|"):
                rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")]); i += 1
            t = doc.add_table(rows=1, cols=len(hdr)); t.style = "Light Grid Accent 1"
            for j, h in enumerate(hdr): add_runs(t.rows[0].cells[j].paragraphs[0], h)
            for row in rows:
                cells = t.add_row().cells
                for j in range(len(hdr)):
                    add_runs(cells[j].paragraphs[0], row[j] if j < len(row) else "")
            for row in t.rows:
                for c in row.cells:
                    for p in c.paragraphs:
                        for r in p.runs: r.font.size = Pt(9)
            continue
        if ln.startswith(">"):
            p = doc.add_paragraph(style="Intense Quote"); add_runs(p, ln.lstrip("> ").strip()); i += 1; continue
        m = re.match(r"^(\s*)[-*]\s+(.*)", ln)
        if m:
            text = m.group(2).replace("[ ] ", "☐ ").replace("[x] ", "☑ ")
            p = doc.add_paragraph(style="List Bullet 2" if m.group(1) else "List Bullet"); add_runs(p, text); i += 1; continue
        m = re.match(r"^\s*(\d+)\.\s+(.*)", ln)
        if m:
            p = doc.add_paragraph(style="List Number"); add_runs(p, m.group(2)); i += 1; continue
        if ln.strip() == "---" or not ln.strip():
            i += 1; continue
        p = doc.add_paragraph(); add_runs(p, ln.strip()); i += 1
    doc.save(DIST / "oshima-tourism-ai-knowledge.docx")

def main():
    DIST.mkdir(exist_ok=True)
    md_text = combined_md()
    (DIST / "oshima-tourism-ai-knowledge.md").write_text(md_text, encoding="utf-8")
    html_path = build_html(md_text)
    build_docx(md_text)
    build_pdf(html_path)
    for f in sorted(DIST.iterdir()): print(f"{f.name}\t{f.stat().st_size//1024} KB")

if __name__ == "__main__":
    main()
